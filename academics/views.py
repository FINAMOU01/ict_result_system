from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db import transaction
from django.db.models import Count, Q, Max
from django.http import JsonResponse
from django.urls import reverse
from .models import Semester, Course, Student, Enrollment, ImportLog, AdmittedStudent
from .forms import SemesterForm, ImportFileForm, AssignProfessorForm
from .utils import import_enrollment_file, build_import_preview, import_enrollment_records
from accounts.models import CustomUser, ActivityLog
from results.models import Grade


def role_required(roles):
    def decorator(view_func):
        def wrapper(request, *args, **kwargs):
            if not request.user.is_authenticated:
                return redirect('login')
            if request.user.role not in roles:
                messages.error(request, "Access denied.")
                return redirect('dashboard')
            return view_func(request, *args, **kwargs)
        wrapper.__name__ = view_func.__name__
        return wrapper
    return decorator


@login_required
def dashboard_redirect(request):
    from accounts.views import dashboard_redirect as acc_redirect
    return acc_redirect(request)


# ─── ADMIN VIEWS ──────────────────────────────────────────────────────────────

@login_required
@role_required(['admin'])
def admin_dashboard(request):
    active_semester = Semester.objects.filter(is_active=True).first()
    all_semesters = Semester.objects.all().order_by('-created_at')
    stats = {
        'semesters': Semester.objects.count(),
        'students': Student.objects.count(),
        'courses': Course.objects.filter(semester=active_semester).count() if active_semester else 0,
        'professors': CustomUser.objects.filter(role='professor', is_active=True).count(),
        'enrollments': Enrollment.objects.filter(semester=active_semester).count() if active_semester else 0,
        'import_logs': ImportLog.objects.order_by('-created_at')[:5],
        'pending_assignments': Course.objects.filter(semester=active_semester, professor__isnull=True).count() if active_semester else 0,
    }
    return render(request, 'admin_panel/dashboard.html', {
        'active_semester': active_semester,
        'all_semesters': all_semesters,
        'stats': stats,
    })


@login_required
@role_required(['admin'])
@login_required
@role_required(['admin'])
def semester_list(request):
    semesters = Semester.objects.annotate(
        course_count=Count('courses', distinct=True),
        enrollment_count=Count('enrollments', distinct=True)
    ).order_by('-created_at')
    form = SemesterForm()
    if request.method == 'POST':
        form = SemesterForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Semester created successfully.")
            return redirect('semester_list')
    return render(request, 'admin_panel/semesters.html', {'semesters': semesters, 'form': form})


@login_required
@role_required(['admin'])
def semester_edit(request, semester_id):
    semester = get_object_or_404(Semester, id=semester_id)
    form = SemesterForm(request.POST or None, instance=semester)

    if request.method == 'POST' and form.is_valid():
        form.save()
        messages.success(request, "Semester updated successfully.")
        return redirect('semester_list')

    return render(request, 'admin_panel/semester_edit.html', {
        'form': form,
        'semester': semester,
    })


@login_required
@role_required(['admin'])
def import_file(request):
    form = ImportFileForm()

    preview_data = None
    if request.session.get('import_preview'):
        preview_data = request.session.get('import_preview')

    if request.method == 'POST':
        if 'preview_import' in request.POST:
            form = ImportFileForm(request.POST, request.FILES)
            if form.is_valid():
                try:
                    preview = build_import_preview(request.FILES['file'])
                    request.session['import_preview'] = {
                        'file_name': preview['file_name'],
                        'semester_id': form.cleaned_data['semester'].id,
                        'records': preview['records'],
                        'course_summary': preview['course_summary'],
                        'course_details': preview['course_details'],
                        'preview_rows': preview['preview_rows'],
                        'total_rows': preview['total_rows'],
                        'total_courses': preview['total_courses'],
                        'errors': preview['errors'],
                    }
                    preview_data = request.session['import_preview']
                    messages.success(request, "Preview generated. Validate to import the data.")
                except ValueError as e:
                    messages.error(request, str(e))
                except Exception as e:
                    messages.error(request, f"Preview failed: {str(e)}")

        elif 'confirm_import' in request.POST:
            preview_data = request.session.get('import_preview')
            if not preview_data:
                messages.error(request, "No preview found. Upload and preview a file first.")
                return redirect('import_file')
            try:
                log = import_enrollment_records(
                    preview_data.get('records', []),
                    preview_data.get('semester_id'),
                    request.user,
                    preview_data.get('file_name', 'imported_file.csv'),
                )
                request.session.pop('import_preview', None)
                messages.success(
                    request,
                    f"Import successful! {log.students_created} new students, "
                    f"{log.courses_created} new courses, {log.enrollments_created} enrollments created."
                )
                if log.errors:
                    messages.warning(request, f"Some rows had errors:\n{log.errors[:500]}")
                return redirect('admin_dashboard')
            except ValueError as e:
                messages.error(request, str(e))
            except Exception as e:
                messages.error(request, f"Import failed: {str(e)}")

        elif 'cancel_preview' in request.POST:
            request.session.pop('import_preview', None)
            messages.info(request, "Import preview cancelled.")
            return redirect('import_file')

    return render(request, 'admin_panel/import_file.html', {
        'form': form,
        'preview_data': preview_data,
    })


@login_required
@role_required(['admin'])
def course_list(request):
    semester_id = request.GET.get('semester')
    level_filter = request.GET.get('level')
    semesters = Semester.objects.all().order_by('-created_at')
    active_semester = Semester.objects.filter(is_active=True).first()
    selected_semester = None

    courses = Course.objects.select_related('professor', 'semester').annotate(
        student_count=Count('enrollments', distinct=True)
    )
    if semester_id:
        courses = courses.filter(semester_id=semester_id)
        selected_semester = get_object_or_404(Semester, id=semester_id)
    elif active_semester:
        courses = courses.filter(semester=active_semester)
        selected_semester = active_semester
    if level_filter:
        courses = courses.filter(level=level_filter)

    return render(request, 'admin_panel/courses.html', {
        'courses': courses.order_by('name'),
        'semesters': semesters,
        'selected_semester': selected_semester,
        'level_filter': level_filter,
    })


@login_required
@role_required(['admin'])
def admin_course_students(request, course_id):
    course = get_object_or_404(
        Course.objects.select_related('semester', 'professor'),
        id=course_id
    )
    query = request.GET.get('q', '').strip()

    enrollments = course.enrollments.select_related('student').order_by('student__last_name', 'student__first_name')
    if query:
        enrollments = enrollments.filter(
            Q(student__matricule__icontains=query) |
            Q(student__last_name__icontains=query) |
            Q(student__first_name__icontains=query) |
            Q(student__email__icontains=query)
        )

    return render(request, 'admin_panel/course_students.html', {
        'course': course,
        'enrollments': enrollments,
        'query': query,
    })


@login_required
@role_required(['admin'])
def assign_professor(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    form = AssignProfessorForm(request.POST or None, instance=course)
    if request.method == 'POST' and form.is_valid():
        form.save()
        prof = course.professor
        msg = f"Professor {prof.get_full_name()} assigned to {course.code}." if prof else f"Professor unassigned from {course.code}."
        messages.success(request, msg)
        return redirect('course_list')
    return render(request, 'admin_panel/assign_professor.html', {'form': form, 'course': course})


@login_required
@role_required(['admin'])
def student_list(request):
    query = request.GET.get('q', '')
    level_filter = request.GET.get('level', '')
    students = Student.objects.annotate(
        enrollment_count=Count('enrollments', distinct=True)
    ).prefetch_related('enrollments__course')
    if query:
        students = students.filter(
            Q(matricule__icontains=query) |
            Q(last_name__icontains=query) |
            Q(first_name__icontains=query)
        )
    if level_filter:
        students = students.filter(level=level_filter)
    return render(request, 'admin_panel/students.html', {
        'students': students.order_by('last_name'),
        'query': query,
        'level_filter': level_filter,
    })


@login_required
@role_required(['admin', 'registra'])
def admissions_registry(request):
    query = request.GET.get('q', '').strip()
    semester_id = request.GET.get('semester', '').strip()

    students = AdmittedStudent.objects.all()
    semesters = Semester.objects.all().order_by('-created_at')
    
    if query:
        students = students.filter(
            Q(matricule__icontains=query) |
            Q(last_name__icontains=query) |
            Q(first_name__icontains=query) |
            Q(email__icontains=query)
        )
    
    if semester_id:
        students = students.filter(semester_id=semester_id)

    return render(request, 'admissions/registry.html', {
        'students': students.order_by('-created_at', 'last_name', 'first_name'),
        'query': query,
        'semesters': semesters,
        'selected_semester_id': semester_id,
    })


# ─── REGISTRA VIEWS ──────────────────────────────────────────────────────────

@login_required
@role_required(['registra'])
def registra_dashboard(request):
    active_semester = Semester.objects.filter(is_active=True).first()

    if active_semester:
        courses = Course.objects.filter(semester=active_semester).annotate(
            student_count=Count('enrollments', distinct=True)
        ).order_by('name')
        courses_coded = courses.filter(is_coded=True).count()
        courses_decoded = courses.filter(is_decoded=True).count()
        total_students = Enrollment.objects.filter(semester=active_semester).values('student').distinct().count()
    else:
        courses = Course.objects.none()
        courses_coded = 0
        courses_decoded = 0
        total_students = 0

    return render(request, 'registra/dashboard.html', {
        'active_semester': active_semester,
        'courses': courses,
        'courses_coded': courses_coded,
        'courses_decoded': courses_decoded,
        'total_students': total_students,
    })


@login_required
@role_required(['registra'])
def registra_semester_history(request):
    from django.db.models import Count, Q
    
    # Fetch all semesters with courses
    query = request.GET.get('q', '').strip()
    
    all_semesters = Semester.objects.filter(
        courses__isnull=False
    ).distinct().order_by('-created_at')
    
    # Filter by search query if provided
    if query:
        all_semesters = all_semesters.filter(
            Q(name__icontains=query) |
            Q(academic_year__icontains=query)
        )
    
    # Collect data for each semester
    semesters_with_data = []
    total_all_courses = 0
    total_all_graded = 0
    total_all_decoded = 0
    
    active_semester = Semester.objects.filter(is_active=True).first()
    
    for semester in all_semesters:
        courses = Course.objects.filter(semester=semester).annotate(
            student_count=Count('enrollments', distinct=True)
        ).order_by('name')
        
        # Count graded and decoded for this semester
        graded_count = courses.filter(grades_submitted=True).count()
        decoded_count = courses.filter(is_decoded=True).count()
        
        semesters_with_data.append({
            'semester': semester,
            'courses': courses,
            'graded_count': graded_count,
            'decoded_count': decoded_count,
        })
        
        # Accumulate totals
        total_all_courses += courses.count()
        total_all_graded += graded_count
        total_all_decoded += decoded_count
    
    return render(request, 'registra/semester_history.html', {
        'active_semester': active_semester,
        'semesters_with_data': semesters_with_data,
        'query': query,
        'total_courses': total_all_courses,
        'total_graded': total_all_graded,
        'total_decoded': total_all_decoded,
    })


def calculate_semester_statistics(semester):
    """Calculate comprehensive statistics for a semester."""
    from collections import defaultdict
    
    # Get all grades for decoded courses in this semester
    grades = Grade.objects.filter(
        enrollment__course__semester=semester,
        enrollment__course__is_decoded=True,
        letter_grade__isnull=False
    ).select_related('enrollment__course')
    
    # Total courses and students in semester
    total_courses = Course.objects.filter(semester=semester).count()
    total_students = Enrollment.objects.filter(
        course__semester=semester
    ).values('student').distinct().count()
    
    # Basic counts
    total_graded = grades.count()
    pass_grades = ['A', 'B+', 'B', 'C+', 'C']
    pass_count = grades.filter(letter_grade__in=pass_grades).count()
    fail_count = total_graded - pass_count
    pass_pct = round((pass_count / total_graded * 100) if total_graded > 0 else 0)
    fail_pct = 100 - pass_pct
    
    # Grade distribution
    grade_counts = defaultdict(int)
    for grade_choice in ['A', 'B+', 'B', 'C+', 'C', 'D+', 'D', 'F']:
        count = grades.filter(letter_grade=grade_choice).count()
        if count > 0:
            grade_counts[grade_choice] = count
    
    # Grade percentages
    grade_pcts = {}
    for grade, count in grade_counts.items():
        grade_pcts[grade] = round((count / total_graded * 100) if total_graded > 0 else 0)
    
    # Per-course statistics
    course_stats = []
    courses = Course.objects.filter(semester=semester, is_decoded=True)
    
    for course in courses:
        course_grades = grades.filter(enrollment__course=course)
        course_total = course_grades.count()
        course_pass = course_grades.filter(letter_grade__in=pass_grades).count()
        course_fail = course_total - course_pass
        course_pass_pct = round((course_pass / course_total * 100) if course_total > 0 else 0)
        
        # Grade distribution for this course
        course_grade_counts = defaultdict(int)
        for grade_choice in ['A', 'B+', 'B', 'C+', 'C', 'D+', 'D', 'F']:
            count = course_grades.filter(letter_grade=grade_choice).count()
            if count > 0:
                course_grade_counts[grade_choice] = count
        
        course_stats.append({
            'course': course,
            'total_graded': course_total,
            'pass_count': course_pass,
            'fail_count': course_fail,
            'pass_pct': course_pass_pct,
            'fail_pct': 100 - course_pass_pct,
            'grade_counts': dict(course_grade_counts),
        })
    
    # Find best and worst courses
    best_course = None
    worst_course = None
    if course_stats:
        best_course = max(course_stats, key=lambda x: x['pass_pct'])
        worst_course = min(course_stats, key=lambda x: x['pass_pct'])
    
    return {
        'total_courses': total_courses,
        'total_students': total_students,
        'total_graded': total_graded,
        'pass_count': pass_count,
        'fail_count': fail_count,
        'pass_pct': pass_pct,
        'fail_pct': fail_pct,
        'grade_counts': dict(grade_counts),
        'grade_pcts': grade_pcts,
        'course_stats': course_stats,
        'best_course': best_course,
        'worst_course': worst_course,
    }


@login_required
@role_required(['registra'])
def registra_semester_report(request, semester_id):
    """Display semester statistics and report."""
    semester = get_object_or_404(Semester, id=semester_id)
    
    # Check if semester has decoded courses
    decoded_courses = Course.objects.filter(semester=semester, is_decoded=True).count()
    if decoded_courses == 0:
        messages.warning(request, "This semester has no decoded courses yet.")
        return redirect('registra_semester_history')
    
    stats = calculate_semester_statistics(semester)
    
    return render(request, 'registra/semester_report.html', {
        'semester': semester,
        'stats': stats,
    })


@login_required
@role_required(['registra'])
def export_report_pdf(request, semester_id):
    """Export semester report as PDF."""
    from io import BytesIO
    from django.http import FileResponse
    from datetime import datetime
    
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
    except ImportError:
        messages.error(request, "PDF library not available. Please install reportlab.")
        return redirect('registra_semester_history')
    
    semester = get_object_or_404(Semester, id=semester_id)
    stats = calculate_semester_statistics(semester)
    
    if stats['total_graded'] == 0:
        messages.warning(request, "No grade data available for this semester.")
        return redirect('registra_semester_history')
    
    # Create PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#1A4A8A'),
        spaceAfter=12,
        alignment=1,
    )
    
    # Add title
    story.append(Paragraph(f"Semester Report: {semester.name}", title_style))
    story.append(Paragraph(f"Academic Year: {semester.academic_year}", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Summary statistics table
    story.append(Paragraph("Summary Statistics", styles['Heading2']))
    summary_data = [
        ['Metric', 'Value', 'Percentage'],
        ['Total Courses', str(stats['total_courses']), '—'],
        ['Total Students', str(stats['total_students']), '—'],
        ['Total Graded', str(stats['total_graded']), '100%'],
        ['Pass Count', str(stats['pass_count']), f"{stats['pass_pct']}%"],
        ['Fail Count', str(stats['fail_count']), f"{stats['fail_pct']}%"],
    ]
    summary_table = Table(summary_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A4A8A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # Grade distribution table
    story.append(Paragraph("Grade Distribution", styles['Heading2']))
    grade_data = [['Grade', 'Count', 'Percentage']]
    for grade in ['A', 'B+', 'B', 'C+', 'C', 'D+', 'D', 'F']:
        if grade in stats['grade_counts']:
            grade_data.append([
                grade,
                str(stats['grade_counts'][grade]),
                f"{stats['grade_pcts'].get(grade, 0)}%"
            ])
    
    if len(grade_data) > 1:
        grade_table = Table(grade_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
        grade_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A4A8A')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
        ]))
        story.append(grade_table)
        story.append(Spacer(1, 20))
    
    # Performance highlights section
    if stats['best_course'] or stats['worst_course']:
        story.append(Paragraph("Performance Highlights", styles['Heading2']))
        perf_data = []
        
        if stats['best_course']:
            perf_data.append(['🏆 Best Course', stats['best_course']['course'].code, f"{stats['best_course']['pass_pct']}%"])
        if stats['worst_course']:
            perf_data.append(['⚠️ Most Difficult', stats['worst_course']['course'].code, f"{stats['worst_course']['pass_pct']}%"])
        
        if perf_data:
            perf_table = Table(perf_data, colWidths=[2*inch, 2*inch, 2*inch])
            perf_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ]))
            story.append(perf_table)
            story.append(Spacer(1, 20))
    
    # Course breakdown section
    if stats['course_stats']:
        story.append(Paragraph("Course Breakdown", styles['Heading2']))
        
        for course_stat in stats['course_stats']:
            story.append(Paragraph(f"{course_stat['course'].code} — {course_stat['course'].name}", styles['Heading3']))
            
            course_data = [
                ['Metric', 'Count', 'Percentage', 'Status'],
                ['Total Graded', str(course_stat['total_graded']), '100%', '—'],
                ['Pass', str(course_stat['pass_count']), f"{course_stat['pass_pct']}%", '✓'],
                ['Fail', str(course_stat['fail_count']), f"{course_stat['fail_pct']}%", '✗'],
            ]
            
            course_table = Table(course_data, colWidths=[2*inch, 1.2*inch, 1.2*inch, 0.8*inch])
            course_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#059669')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#E0F5E9')),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))
            story.append(course_table)
            story.append(Spacer(1, 10))
    
    # Build PDF
    try:
        doc.build(story)
        buffer.seek(0)
        filename = f"Semester_Report_{semester.name}_{datetime.now().strftime('%Y%m%d')}.pdf"
        return FileResponse(buffer, as_attachment=True, filename=filename)
    except Exception as e:
        messages.error(request, f"Error generating PDF: {str(e)}")
        return redirect('registra_semester_history')


@login_required
@role_required(['registra'])
def export_report_docx(request, semester_id):
    """Export semester report as Word document."""
    from io import BytesIO
    from django.http import FileResponse
    from datetime import datetime
    
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        messages.error(request, "Word library not available. Please install python-docx.")
        return redirect('registra_semester_history')
    
    semester = get_object_or_404(Semester, id=semester_id)
    stats = calculate_semester_statistics(semester)
    
    if stats['total_graded'] == 0:
        messages.warning(request, "No grade data available for this semester.")
        return redirect('registra_semester_history')
    
    try:
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f"Semester Report: {semester.name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Academic Year: {semester.academic_year}\n").bold = True
        meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_paragraph()
        
        # Summary section
        doc.add_heading("Summary Statistics", level=1)
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Light Grid Accent 1'
        
        headers = table.rows[0].cells
        headers[0].text = "Metric"
        headers[1].text = "Value"
        headers[2].text = "Percentage"
        
        data_rows = [
            ["Total Courses", str(stats['total_courses']), "—"],
            ["Total Students", str(stats['total_students']), "—"],
            ["Total Graded", str(stats['total_graded']), "100%"],
            ["Pass Count", str(stats['pass_count']), f"{stats['pass_pct']}%"],
            ["Fail Count", str(stats['fail_count']), f"{stats['fail_pct']}%"],
        ]
        
        for i, (metric, value, pct) in enumerate(data_rows, 1):
            cells = table.rows[i].cells
            cells[0].text = metric
            cells[1].text = value
            cells[2].text = pct
        
        doc.add_paragraph()
        
        # Grade distribution section
        doc.add_heading("Grade Distribution", level=1)
        grade_table = doc.add_table(rows=1, cols=3)
        grade_table.style = 'Light Grid Accent 1'
        
        grade_headers = grade_table.rows[0].cells
        grade_headers[0].text = "Grade"
        grade_headers[1].text = "Count"
        grade_headers[2].text = "Percentage"
        
        for grade in ['A', 'B+', 'B', 'C+', 'C', 'D+', 'D', 'F']:
            if grade in stats['grade_counts']:
                row = grade_table.add_row().cells
                row[0].text = grade
                row[1].text = str(stats['grade_counts'][grade])
                row[2].text = f"{stats['grade_pcts'].get(grade, 0)}%"
        
        doc.add_paragraph()
        
        # Performance highlights section
        if stats['best_course'] or stats['worst_course']:
            doc.add_heading("Performance Highlights", level=1)
            perf_table = doc.add_table(rows=1, cols=3)
            perf_table.style = 'Light Grid Accent 1'
            
            p_headers = perf_table.rows[0].cells
            p_headers[0].text = "Status"
            p_headers[1].text = "Course Code"
            p_headers[2].text = "Pass Rate"
            
            if stats['best_course']:
                row = perf_table.add_row().cells
                row[0].text = "🏆 Best Course"
                row[1].text = stats['best_course']['course'].code
                row[2].text = f"{stats['best_course']['pass_pct']}%"
            
            if stats['worst_course']:
                row = perf_table.add_row().cells
                row[0].text = "⚠️ Most Difficult"
                row[1].text = stats['worst_course']['course'].code
                row[2].text = f"{stats['worst_course']['pass_pct']}%"
            
            doc.add_paragraph()
        
        # Course breakdown section
        if stats['course_stats']:
            doc.add_heading("Course Breakdown", level=1)
            
            for course_stat in stats['course_stats']:
                doc.add_heading(f"{course_stat['course'].code} — {course_stat['course'].name}", level=2)
                
                # Main metrics table
                course_table = doc.add_table(rows=1, cols=4)
                course_table.style = 'Light Grid Accent 1'
                
                c_headers = course_table.rows[0].cells
                c_headers[0].text = "Metric"
                c_headers[1].text = "Count"
                c_headers[2].text = "Percentage"
                c_headers[3].text = "Status"
                
                metrics = [
                    ["Total Graded", str(course_stat['total_graded']), "100%", "—"],
                    ["Pass", str(course_stat['pass_count']), f"{course_stat['pass_pct']}%", "✓"],
                    ["Fail", str(course_stat['fail_count']), f"{course_stat['fail_pct']}%", "✗"],
                ]
                
                for metric, count, pct, status in metrics:
                    row = course_table.add_row().cells
                    row[0].text = metric
                    row[1].text = count
                    row[2].text = pct
                    row[3].text = status
                
                # Grade distribution for this course
                if course_stat['grade_counts']:
                    doc.add_paragraph()
                    doc.add_paragraph("Grade Distribution:").runs[0].bold = True
                    
                    grade_dist_table = doc.add_table(rows=1, cols=3)
                    grade_dist_table.style = 'Light Grid Accent 1'
                    
                    g_headers = grade_dist_table.rows[0].cells
                    g_headers[0].text = "Grade"
                    g_headers[1].text = "Count"
                    g_headers[2].text = "Percentage"
                    
                    total_course_grades = sum(course_stat['grade_counts'].values())
                    for grade in ['A', 'B+', 'B', 'C+', 'C', 'D+', 'D', 'F']:
                        if grade in course_stat['grade_counts']:
                            count = course_stat['grade_counts'][grade]
                            pct = round((count / total_course_grades * 100) if total_course_grades > 0 else 0)
                            row = grade_dist_table.add_row().cells
                            row[0].text = grade
                            row[1].text = str(count)
                            row[2].text = f"{pct}%"
                
                doc.add_paragraph()
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"Semester_Report_{semester.name}_{datetime.now().strftime('%Y%m%d')}.docx"
        return FileResponse(buffer, as_attachment=True, filename=filename)
        
    except Exception as e:
        messages.error(request, f"Error generating Word document: {str(e)}")
        return redirect('registra_semester_history')


@login_required
@role_required(['registra'])
def registra_course_detail(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    
    # Order by anonymous_code for coded courses, by student name otherwise
    if course.is_coded:
        enrollments = course.enrollments.select_related('student').order_by('anonymous_code')
    else:
        enrollments = course.enrollments.select_related('student').order_by('student__last_name')
    
    query = request.GET.get('q', '')
    if query:
        enrollments = enrollments.filter(
            Q(student__matricule__icontains=query) |
            Q(student__last_name__icontains=query) |
            Q(student__first_name__icontains=query)
        )
    return render(request, 'registra/course_detail.html', {
        'course': course,
        'enrollments': enrollments,
        'query': query,
    })


@login_required
@role_required(['registra'])
def add_walkin_student(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    
    # Check if course is already coded
    if course.is_coded:
        messages.error(request, "Cannot add students after coding has started.")
        return redirect('registra_course_detail', course_id=course_id)
    
    if request.method == 'POST':
        from .forms import WalkInStudentForm
        form = WalkInStudentForm(request.POST)
        if form.is_valid():
            matricule = form.cleaned_data['matricule'].strip()
            first_name = form.cleaned_data['first_name'].strip().title()
            last_name = form.cleaned_data['last_name'].strip().upper()
            
            # Check if student exists
            try:
                student = Student.objects.get(matricule=matricule)
            except Student.DoesNotExist:
                # Create new walk-in student
                student = Student.objects.create(
                    matricule=matricule,
                    first_name=first_name,
                    last_name=last_name,
                    level='bachelor',  # Default level for walk-ins
                    is_walkin=True
                )
            
            # Check if already enrolled in this course
            enrollment_exists = Enrollment.objects.filter(
                student=student,
                course=course,
                semester=course.semester
            ).exists()
            
            if enrollment_exists:
                messages.error(request, f"This student ({student.matricule}) is already enrolled in this course.")
                return redirect('add_walkin_student', course_id=course_id)
            
            # Create enrollment
            enrollment = Enrollment.objects.create(
                student=student,
                course=course,
                semester=course.semester
            )
            
            # Create associated Grade object
            from results.models import Grade
            Grade.objects.get_or_create(enrollment=enrollment)
            
            messages.success(request, f"✅ {student.full_name()} ({student.matricule}) added to {course.code}.")
            return redirect('registra_course_detail', course_id=course_id)
    else:
        from .forms import WalkInStudentForm
        form = WalkInStudentForm()
    
    return render(request, 'registra/add_walkin_student.html', {
        'form': form,
        'course': course,
    })


@login_required
@role_required(['registra'])
def add_from_admissions(request, course_id):
    """Add a student from admitted students list to a course."""
    course = get_object_or_404(Course, id=course_id)

    if course.is_coded:
        messages.error(request, "Cannot add students after coding has started.")
        return redirect('registra_course_detail', course_id=course_id)

    if request.method == 'POST':
        matricule = request.POST.get('matricule', '').strip()

        if not matricule:
            messages.error(request, "Please enter a matricule number.")
            return redirect('add_from_admissions', course_id=course_id)

        try:
            # Search across all semesters - get the most recent record
            admitted_student = AdmittedStudent.objects.filter(
                matricule=matricule
            ).order_by('-semester__created_at').first()
            
            if not admitted_student:
                messages.error(request, f"No admitted student found with matricule '{matricule}'.")
                return redirect('add_from_admissions', course_id=course_id)
        except AdmittedStudent.DoesNotExist:
            messages.error(request, f"No admitted student found with matricule '{matricule}'.")
            return redirect('add_from_admissions', course_id=course_id)

        enrollment_exists = Enrollment.objects.filter(
            student__matricule=matricule,
            course=course,
            semester=course.semester
        ).exists()
        if enrollment_exists:
            messages.error(request, f"This student ({matricule}) is already enrolled in this course.")
            return redirect('add_from_admissions', course_id=course_id)

        student, _ = Student.objects.get_or_create(
            matricule=matricule,
            defaults={
                'first_name': admitted_student.first_name,
                'last_name': admitted_student.last_name,
                'email': admitted_student.email or '',
                'level': admitted_student.level or 'bachelor',
                'is_walkin': False,
            }
        )

        Enrollment.objects.create(
            student=student,
            course=course,
            semester=course.semester
        )
        
        # Create associated Grade object
        from results.models import Grade
        enrollment = Enrollment.objects.get(student=student, course=course)
        Grade.objects.get_or_create(enrollment=enrollment)

        messages.success(request, f"{student.full_name()} ({student.matricule}) added to {course.code}.")
        return redirect('registra_course_detail', course_id=course_id)

    admitted_query = request.GET.get('q', '').strip()
    admitted_students = AdmittedStudent.objects.all()
    if admitted_query:
        admitted_students = admitted_students.filter(
            Q(matricule__icontains=admitted_query) |
            Q(last_name__icontains=admitted_query) |
            Q(first_name__icontains=admitted_query)
        )

    return render(request, 'registra/add_from_admissions.html', {
        'course': course,
        'admitted_students': admitted_students[:80],
        'query': admitted_query,
    })


@login_required
@role_required(['registra'])
def search_admitted_students(request, course_id):
    """AJAX endpoint to search admitted students."""
    try:
        if request.method != 'POST':
            return JsonResponse({'error': 'Method not allowed'}, status=405)
        
        import json
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        
        query = data.get('q', '').strip()
        
        if not query:
            return JsonResponse({'students': []})
        
        course = get_object_or_404(Course, id=course_id)
        
        admitted_students = AdmittedStudent.objects.filter(
            Q(matricule__icontains=query) |
            Q(last_name__icontains=query) |
            Q(first_name__icontains=query)
        ).order_by('matricule')[:50]
        
        # Check which ones are already enrolled
        enrolled_matricules = set(
            Enrollment.objects.filter(course=course).values_list(
                'student__matricule', flat=True
            )
        )
        
        students = []
        for student in admitted_students:
            students.append({
                'matricule': student.matricule,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'level': student.get_level_display() if student.level else 'N/A',
                'admitted_year': student.semester.name if student.semester else 'N/A',
                'already_enrolled': student.matricule in enrolled_matricules
            })
        
        return JsonResponse({'students': students})
    
    except Exception as e:
        import traceback
        print(f"Error in search_admitted_students: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({'error': str(e), 'students': []}, status=500)


@login_required
@role_required(['registra'])
def add_from_admissions_ajax(request, course_id):
    """AJAX endpoint to add student from admitted list."""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Method not allowed'}, status=405)
    
    import json
    data = json.loads(request.body)
    matricule = data.get('matricule', '').strip()
    
    course = get_object_or_404(Course, id=course_id)
    
    if not matricule:
        return JsonResponse({
            'success': False,
            'error': 'Matricule is required'
        })
    
    try:
        # Check if student is already enrolled in this course
        existing_enrollment = Enrollment.objects.filter(
            student__matricule=matricule,
            course=course,
            semester=course.semester
        ).first()
        
        if existing_enrollment:
            return JsonResponse({
                'success': False,
                'error': f'This student is already enrolled in this course'
            })
        
        # Find admitted student
        admitted_student = AdmittedStudent.objects.filter(
            matricule=matricule
        ).order_by('-semester__created_at').first()
        
        if not admitted_student:
            return JsonResponse({
                'success': False,
                'error': f'No admitted student found with matricule "{matricule}"'
            })
        
        # Get or create student
        student, _ = Student.objects.get_or_create(
            matricule=matricule,
            defaults={
                'first_name': admitted_student.first_name,
                'last_name': admitted_student.last_name,
                'email': admitted_student.email or '',
                'level': admitted_student.level or 'bachelor',
                'is_walkin': False,
            }
        )
        
        # Generate anonymous code for coded courses
        anonymous_code = None
        if course.is_coded:
            from django.db.models import Max
            max_code = course.enrollments.aggregate(Max('anonymous_code'))['anonymous_code__max'] or 0
            anonymous_code = max_code + 1
        
        # Create enrollment
        enrollment = Enrollment.objects.create(
            student=student,
            course=course,
            semester=course.semester,
            anonymous_code=anonymous_code
        )
        
        # Create grade record
        from results.models import Grade
        Grade.objects.get_or_create(enrollment=enrollment)
        
        return JsonResponse({
            'success': True,
            'message': f'{student.full_name()} added to {course.code}',
            'student_name': student.full_name(),
            'anonymous_code': anonymous_code or 'N/A'
        })
    
    except Exception as e:
        import traceback
        print(f"Error in add_from_admissions_ajax: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@role_required(['registra'])
def generate_codes(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    if course.is_coded:
        messages.warning(request, f"Codes already generated for {course.code}. Cannot re-code.")
        return redirect('registra_course_detail', course_id=course_id)
    if request.method == 'POST':
        from results.utils import generate_codes_for_course
        count = generate_codes_for_course(course)
        messages.success(request, f"✅ {count} anonymous codes generated for {course.code} - {course.name}.")
        return redirect('registra_course_detail', course_id=course_id)
    return render(request, 'registra/confirm_code.html', {'course': course})


@login_required
@role_required(['registra'])
def download_coding_sheet(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    if not course.is_coded:
        messages.error(request, "Generate codes first before downloading.")
        return redirect('registra_course_detail', course_id=course_id)
    from results.utils import generate_coding_sheet
    from django.http import HttpResponse
    buffer = generate_coding_sheet(course)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="CONFIDENTIAL_{course.code}_coding_sheet.xlsx"'
    return response


@login_required
@role_required(['registra'])
def lookup_code(request):
    """Registra can look up a matricule to find the student's code."""
    result = None
    active_semester = Semester.objects.filter(is_active=True).first()
    if request.method == 'GET' and request.GET.get('matricule'):
        matricule = request.GET.get('matricule', '').strip()
        course_id = request.GET.get('course_id')
        try:
            if course_id:
                enrollment = Enrollment.objects.select_related('student', 'course').get(
                    student__matricule__iexact=matricule,
                    course_id=course_id
                )
            else:
                enrollment = Enrollment.objects.select_related('student', 'course').filter(
                    student__matricule__iexact=matricule,
                    course__semester=active_semester
                ).first()
            result = enrollment
        except Enrollment.DoesNotExist:
            messages.error(request, f"No enrollment found for matricule '{matricule}'.")
    courses = Course.objects.filter(semester=active_semester, is_coded=True).order_by('name') if active_semester else []
    return render(request, 'registra/lookup_code.html', {
        'result': result,
        'courses': courses,
        'active_semester': active_semester,
    })


@login_required
@role_required(['registra'])
def download_decoded_results(request, course_id):
    course = get_object_or_404(Course, id=course_id)
    if not course.grades_submitted:
        messages.error(request, "Professor has not submitted grades yet.")
        return redirect('registra_dashboard')
    from results.utils import generate_results_csv
    from django.http import HttpResponse
    buffer = generate_results_csv(course)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='text/csv; charset=utf-8'
    )
    response['Content-Disposition'] = f'attachment; filename="{course.code}_results_{course.semester.name}.csv"'
    return response


# ─── PROFESSOR VIEWS ──────────────────────────────────────────────────────────

@login_required
@role_required(['professor'])
def professor_dashboard(request):
    active_semester = Semester.objects.filter(is_active=True).first()
    if active_semester:
        courses = Course.objects.filter(
            professor=request.user,
            semester=active_semester
        ).annotate(student_count=Count('enrollments', distinct=True)).order_by('name')
    else:
        courses = Course.objects.none()
    return render(request, 'professor/dashboard.html', {
        'active_semester': active_semester,
        'courses': courses,
    })


@login_required
@role_required(['professor'])
def professor_grade_course(request, course_id):
    course = get_object_or_404(Course, id=course_id, professor=request.user)
    if not course.is_coded:
        messages.warning(request, "The registra has not coded this course yet. Please wait.")
        return redirect('professor_dashboard')

    from results.models import Grade
    enrollments = course.enrollments.select_related('grade').order_by('anonymous_code')

    if request.method == 'POST':
        if course.grades_submitted:
            messages.error(request, "Grades already submitted. Cannot modify.")
            return redirect('professor_dashboard')
        errors = []
        for enrollment in enrollments:
            cc_key = f'cc_{enrollment.id}'
            sn_key = f'sn_{enrollment.id}'
            attendance_key = f'attendance_{enrollment.id}'
            
            cc_val = request.POST.get(cc_key, '').strip()
            sn_val = request.POST.get(sn_key, '').strip()
            attendance_val = request.POST.get(attendance_key, '').strip()
            
            try:
                grade, _ = Grade.objects.get_or_create(enrollment=enrollment)
                
                # Validate and set CC score (max 20)
                if cc_val:
                    cc = float(cc_val)
                    if cc > 20:
                        raise ValueError(f"CC score cannot exceed 20 (entered: {cc})")
                    if not 0 <= cc <= 20:
                        raise ValueError("CC score must be between 0 and 20")
                    grade.cc_score = cc
                
                # Validate and set Attendance score (max 10)
                if attendance_val:
                    attendance = float(attendance_val)
                    if attendance > 10:
                        raise ValueError(f"Attendance score cannot exceed 10 (entered: {attendance})")
                    if not 0 <= attendance <= 10:
                        raise ValueError("Attendance score must be between 0 and 10")
                    grade.attendance_score = attendance
                else:
                    # Default attendance to 0 if not provided
                    grade.attendance_score = 0
                
                # Validate and set SN score (max 70)
                if sn_val:
                    sn = float(sn_val)
                    if sn > 70:
                        raise ValueError(f"SN score cannot exceed 70 (entered: {sn})")
                    if not 0 <= sn <= 70:
                        raise ValueError("SN score must be between 0 and 70")
                    grade.sn_score = sn
                
                # Calculate final score and letter grade
                grade.calculate_final()
                grade.save()
            except ValueError as e:
                errors.append(f"Code {enrollment.anonymous_code}: {e}")

        if 'submit_final' in request.POST and not errors:
            from django.utils import timezone
            Grade.objects.filter(enrollment__course=course).update(
                status='submitted',
                submitted_by=request.user,
                submitted_at=timezone.now()
            )
            course.grades_submitted = True
            course.grades_submitted_at = timezone.now()
            course.save()
            messages.success(request, f"✅ Grades for {course.code} submitted to Registra successfully!")
            return redirect('professor_dashboard')
        elif errors:
            for e in errors:
                messages.error(request, e)
        else:
            messages.success(request, "Grades saved. Review and click 'Submit to Registra' when ready.")

    # Determine which tab to show based on submission status
    active_tab = 'cc' if not course.cc_submitted else 'sn'
    
    return render(request, 'professor/grade_course.html', {
        'course': course,
        'enrollments': enrollments,
        'active_tab': active_tab,
    })


@login_required
@role_required(['professor'])
def download_coded_list(request, course_id):
    """Download the coded list with empty Final Score column for lecturer to fill."""
    course = get_object_or_404(Course, id=course_id, professor=request.user)
    
    from django.http import HttpResponse
    from io import BytesIO
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    
    # Get all enrollments ordered by code
    enrollments = course.enrollments.select_related('student').order_by('anonymous_code')
    
    # Create workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Final Scores"
    
    # Add headers
    headers = ['Anonymous Code', 'Final Score /70']
    ws.append(headers)
    
    # Style the header row
    header_fill = PatternFill(start_color='1A4A8A', end_color='1A4A8A', fill_type='solid')
    header_font = Font(color='FFFFFF', bold=True)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add each code with empty score
    for enrollment in enrollments:
        ws.append([enrollment.anonymous_code, ''])
    
    # Set column widths
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 18
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    # Create response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{course.code}_Final_Scores.xlsx"'
    
    return response


@login_required
@role_required(['professor'])
def submit_cc_scores(request, course_id):
    """Submit CA scores for a course."""
    course = get_object_or_404(Course, id=course_id, professor=request.user)
    
    if request.method != 'POST':
        return redirect('professor_grade_course', course_id=course_id)
    
    if course.cc_submitted:
        messages.error(request, "CA scores already submitted. Cannot modify.")
        return redirect('professor_dashboard')
    
    from results.models import Grade
    from django.utils import timezone
    
    enrollments = course.enrollments.select_related('grade')
    errors = []
    
    for enrollment in enrollments:
        cc_key = f'cc_{enrollment.id}'
        cc_val = request.POST.get(cc_key, '').strip()
        
        if not cc_val:
            continue
        
        try:
            grade, _ = Grade.objects.get_or_create(enrollment=enrollment)
            cc = float(cc_val)
            
            if not 0 <= cc <= 30:
                errors.append(f"{enrollment.student.matricule}: CA score must be between 0 and 30")
                continue
            
            grade.cc_score = cc
            grade.cc_submitted = True  # Mark individual grade as submitted
            grade.calculate_final()  # Calculate final score and letter grade
            grade.save()
        except ValueError:
            errors.append(f"{enrollment.student.matricule}: Invalid CA score")
    
    if not errors:
        # Only mark course as CC submitted if no errors
        course.cc_submitted = True
        course.cc_submitted_at = timezone.now()
        course.save()
        messages.success(request, "✅ CA scores submitted successfully!")
    else:
        for error in errors:
            messages.error(request, error)
    
    return redirect('professor_grade_course', course_id=course_id)





@login_required
@role_required(['professor'])
def download_ca_template(request):
    """Download a blank CA marks template for lecturer to fill."""
    from django.http import HttpResponse
    from io import BytesIO
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    
    # Create workbook with a blank template
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "CA Scores"
    
    # Add headers
    headers = ['Matricule', 'Full Name', 'Email', 'CA Score /30']
    ws.append(headers)
    
    # Style the header row
    header_fill = PatternFill(start_color='1A4A8A', end_color='1A4A8A', fill_type='solid')
    header_font = Font(color='FFFFFF', bold=True)
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add empty rows for user to fill in
    for i in range(10):
        ws.append(['', '', '', ''])
    
    # Set column widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 15
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    # Create response
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="CA_Marks_Template.xlsx"'
    
    return response


@login_required
@role_required(['professor'])
def import_ca_marks(request, course_id):
    """Import CA marks from uploaded Excel file."""
    course = get_object_or_404(Course, id=course_id, professor=request.user)
    
    if request.method != 'POST' or not request.FILES.get('ca_file'):
        messages.error(request, "No file uploaded.")
        return redirect('professor_grade_course', course_id=course_id)
    
    try:
        import openpyxl
        from decimal import Decimal, InvalidOperation
        
        ca_file = request.FILES['ca_file']
        
        # Read Excel file
        wb = openpyxl.load_workbook(ca_file)
        ws = wb.active
        
        # Get all enrollments for this course
        matricule_to_enrollment = {e.student.matricule: e for e in course.enrollments.all() if e.student}
        
        # Parse Excel data
        matched = 0
        invalid_scores = []
        not_found = []
        
        # Skip header row, start from row 2
        for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not row or not row[0]:
                continue
            
            matricule = row[0]
            ca_score_raw = row[3] if len(row) > 3 else ""
            
            if not matricule:
                continue
            
            # Try to find enrollment by matricule
            try:
                matricule_str = str(matricule).strip()
                enrollment = matricule_to_enrollment.get(matricule_str)
            except (ValueError, TypeError):
                enrollment = None
            
            if not enrollment:
                not_found.append(f"Matricule {matricule} not found in course")
                continue
            
            # Skip empty scores
            if ca_score_raw == "" or ca_score_raw is None:
                continue
            
            # Validate and save score
            try:
                score = Decimal(str(ca_score_raw)).quantize(Decimal('0.01'))
                
                if not (0 <= score <= 30):
                    invalid_scores.append(f"Matricule {matricule} - Score {ca_score_raw} out of range (0-30)")
                    continue
                
                # Save grade
                from results.models import Grade
                grade, _ = Grade.objects.get_or_create(enrollment=enrollment)
                grade.cc_score = score
                grade.save()
                matched += 1
                
            except (InvalidOperation, TypeError, ValueError) as e:
                invalid_scores.append(f"Matricule {matricule} - Invalid score: {ca_score_raw}")
                continue
        
        # Build result messages
        result_messages = []
        
        if matched > 0:
            result_messages.append(f"✅ {matched} CA score(s) successfully saved for this course.")
        
        if invalid_scores:
            result_messages.append(f"❌ {len(invalid_scores)} invalid score(s) skipped")
            for error in invalid_scores[:5]:  # Show first 5 errors
                result_messages.append(f"  • {error}")
        
        if not_found:
            result_messages.append(f"⚠️ {len(not_found)} matricule(s) not found in course")
        
        full_message = "<br>".join(result_messages)
        
        if matched > 0:
            messages.success(request, full_message, extra_tags='safe')
        else:
            messages.warning(request, full_message if result_messages else "No valid records to import.", extra_tags='safe')
        
        # Redirect to grade course page with cc tab
        url = reverse('professor_grade_course', args=[course_id]) + '?tab=cc'
        return redirect(url)
        
    except Exception as e:
        messages.error(request, f"Error reading file: {str(e)}")
        return redirect('professor_grade_course', course_id=course_id)


@login_required
@role_required(['professor'])
def submit_sn_scores(request, course_id):
    """Submit Final (SN) scores for a course."""
    course = get_object_or_404(Course, id=course_id, professor=request.user)
    
    if request.method != 'POST':
        return redirect('professor_grade_course', course_id=course_id)
    
    if course.sn_submitted:
        messages.error(request, "Final scores already submitted. Cannot modify.")
        return redirect('professor_dashboard')
    
    from results.models import Grade
    from django.utils import timezone
    
    enrollments = course.enrollments.select_related('grade')
    errors = []
    
    for enrollment in enrollments:
        sn_key = f'sn_{enrollment.id}'
        sn_val = request.POST.get(sn_key, '').strip()
        
        if not sn_val:
            continue
        
        try:
            grade, _ = Grade.objects.get_or_create(enrollment=enrollment)
            sn = float(sn_val)
            
            if not 0 <= sn <= 70:
                errors.append(f"{enrollment.anonymous_code}: Final score must be between 0 and 70")
                continue
            
            grade.sn_score = sn
            grade.sn_submitted = True  # Mark individual grade as submitted
            grade.calculate_final()  # Calculate final score and letter grade
            grade.save()
        except ValueError:
            errors.append(f"{enrollment.anonymous_code}: Invalid Final score")
    
    if not errors:
        # Only mark course as SN submitted if no errors
        # Mark all grades as submitted to registra
        Grade.objects.filter(enrollment__course=course).update(
            status='submitted',
            submitted_by=request.user,
            submitted_at=timezone.now()
        )
        course.sn_submitted = True
        course.grades_submitted = True  # Also mark overall submission
        course.grades_submitted_at = timezone.now()
        course.save()
        messages.success(request, f"✅ Final scores for {course.code} submitted to Registra successfully!")
    else:
        for error in errors:
            messages.error(request, error)
        messages.success(request, "✅ Final scores submitted with some errors (see above)!")
    
    return redirect('professor_dashboard')


@login_required
@role_required(['professor'])
def import_sn_marks(request, course_id):
    """Import Final (SN) marks from uploaded Excel file."""
    course = get_object_or_404(Course, id=course_id, professor=request.user)
    
    if request.method != 'POST' or not request.FILES.get('sn_file'):
        messages.error(request, "No file uploaded.")
        return redirect('professor_grade_course', course_id=course_id)
    
    try:
        import openpyxl
        from decimal import Decimal, InvalidOperation
        
        sn_file = request.FILES['sn_file']
        
        # Read Excel file
        wb = openpyxl.load_workbook(sn_file)
        ws = wb.active
        
        # Get all enrollments for this course
        code_to_enrollment = {e.anonymous_code: e for e in course.enrollments.all() if e.anonymous_code}
        
        # Parse Excel data
        matched = 0
        invalid_scores = []
        not_found = []
        
        # Skip header row, start from row 2
        for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not row or not row[0]:
                continue
            
            code = row[0]
            sn_score_raw = row[1] if len(row) > 1 else ""
            
            if not code:
                continue
            
            # Try to find enrollment by code
            try:
                code_int = int(code)
                enrollment = code_to_enrollment.get(code_int)
            except (ValueError, TypeError):
                enrollment = None
            
            if not enrollment:
                not_found.append(f"Code {code} not found in course")
                continue
            
            # Skip empty scores
            if sn_score_raw == "" or sn_score_raw is None:
                continue
            
            # Validate and save score
            try:
                score = Decimal(str(sn_score_raw)).quantize(Decimal('0.01'))
                
                if not (0 <= score <= 70):
                    invalid_scores.append(f"Code {code} - Score {sn_score_raw} out of range (0-70)")
                    continue
                
                # Save grade
                from results.models import Grade
                grade, _ = Grade.objects.get_or_create(enrollment=enrollment)
                grade.sn_score = score
                grade.save()
                matched += 1
                
            except (InvalidOperation, TypeError, ValueError) as e:
                invalid_scores.append(f"Code {code} - Invalid score: {sn_score_raw}")
                continue
        
        # Build result messages
        result_messages = []
        
        if matched > 0:
            result_messages.append(f"✅ {matched} Final score(s) successfully saved for this course.")
        
        if invalid_scores:
            result_messages.append(f"❌ {len(invalid_scores)} invalid score(s) skipped")
            for error in invalid_scores[:5]:  # Show first 5 errors
                result_messages.append(f"  • {error}")
        
        if not_found:
            result_messages.append(f"⚠️ {len(not_found)} code(s) not found in course")
        
        full_message = "<br>".join(result_messages)
        
        if matched > 0:
            messages.success(request, full_message, extra_tags='safe')
        else:
            messages.warning(request, full_message if result_messages else "No valid records to import.", extra_tags='safe')
        
        # Redirect to grade course page with tab parameter
        url = reverse('professor_grade_course', args=[course_id]) + '?tab=sn'
        return redirect(url)
        
    except Exception as e:
        messages.error(request, f"Error reading file: {str(e)}")
        return redirect('professor_grade_course', course_id=course_id)


# ─── REGISTRA VIEWS ───────────────────────────────────────────────────────────

@login_required
@role_required(['registra'])
def registra_view_anonymous_results(request, course_id):
    """View anonymous results (code, CC, SN, Final) before decoding."""
    course = get_object_or_404(Course, id=course_id)
    
    if not course.grades_submitted:
        messages.error(request, "Grades have not been submitted for this course yet.")
        return redirect('registra_dashboard')
    
    enrollments = course.enrollments.select_related('grade').order_by('anonymous_code')
    
    return render(request, 'registra/anonymous_results.html', {
        'course': course,
        'enrollments': enrollments,
    })


@login_required
@role_required(['registra'])
def registra_decode_results(request, course_id):
    """Decode results (POST only)."""
    course = get_object_or_404(Course, id=course_id)
    
    if not course.grades_submitted:
        messages.error(request, "Grades have not been submitted for this course yet.")
        return redirect('registra_dashboard')
    
    course.is_decoded = True
    course.save()
    
    messages.success(request, f"✅ Results for {course.code} have been decoded successfully!")
    return redirect('registra_decoded_results', course_id=course.id)


@login_required
@role_required(['registra'])
def registra_decoded_results(request, course_id):
    """View decoded results with student information."""
    course = get_object_or_404(Course, id=course_id)
    
    if not course.is_decoded:
        messages.error(request, "Results have not been decoded for this course yet.")
        return redirect('registra_dashboard')
    
    enrollments = course.enrollments.select_related('student', 'grade').order_by('student__last_name', 'student__first_name')
    
    # Add remarks to each enrollment
    for enrollment in enrollments:
        if enrollment.grade and enrollment.grade.final_score is not None:
            enrollment.remarks = 'PASS' if enrollment.grade.final_score >= 10 else 'FAIL'
        else:
            enrollment.remarks = 'N/A'
    
    return render(request, 'registra/decoded_results.html', {
        'course': course,
        'enrollments': enrollments,
    })


# ─── ADMIN ACTIVITY LOG VIEW ──────────────────────────────────────────────────

@login_required
@role_required(['admin'])
def activity_log(request):
    """
    Display comprehensive audit log of all system activities for admin review.
    Shows who did what, when, and their IP address.
    """
    # Get filters from request
    action_filter = request.GET.get('action', '')
    user_filter = request.GET.get('user', '')
    query = request.GET.get('q', '').strip()
    status_filter = request.GET.get('status', '')

    # Start with all logs, ordered by most recent first
    logs = ActivityLog.objects.all().order_by('-timestamp').select_related('user')

    # Apply filters
    if action_filter:
        logs = logs.filter(action=action_filter)

    if user_filter:
        logs = logs.filter(user__id=user_filter)

    if query:
        logs = logs.filter(
            Q(description__icontains=query) |
            Q(affected_entity__icontains=query) |
            Q(user__username__icontains=query) |
            Q(user__email__icontains=query)
        )

    if status_filter:
        logs = logs.filter(status=status_filter)

    # Get all users for dropdown filter
    users = CustomUser.objects.filter(is_active=True).order_by('first_name', 'last_name')

    # Paginate logs (50 per page)
    from django.core.paginator import Paginator
    paginator = Paginator(logs, 50)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)

    # Get action choices for filter
    action_choices = ActivityLog._meta.get_field('action').choices

    context = {
        'page_obj': page_obj,
        'logs': page_obj.object_list,
        'action_filter': action_filter,
        'user_filter': user_filter,
        'query': query,
        'status_filter': status_filter,
        'users': users,
        'action_choices': action_choices,
        'total_logs': paginator.count,
    }

    return render(request, 'admin_panel/activity_log.html', context)


# ─── ADMITTED STUDENTS IMPORT & LOOKUP ────────────────────────────────────────

@login_required
@role_required(['admin'])
def select_semester_for_import(request):
    """
    Admin selects a semester to import admitted students for.
    Lists all semesters with import links.
    """
    semesters = Semester.objects.all().order_by('-created_at')
    return render(request, 'admin_panel/select_semester_import.html', {
        'semesters': semesters,
    })


@login_required
@role_required(['admin'])
def import_admitted_students(request, semester_id):
    """
    Admin imports the full list of admitted students for a semester.
    Accepts CSV or Excel (.xlsx/.xls).
    Expected columns: matricule, first_name, last_name, email (optional),
                      program (optional), level (optional).
    No duplicates: uses update_or_create on (semester, matricule).
    """
    semester = get_object_or_404(Semester, id=semester_id)
    
    if request.method == 'GET':
        return render(request, 'admin_panel/import_admitted_students.html', {
            'semester': semester,
        })
    
    if request.method == 'POST':
        if 'file' not in request.FILES:
            messages.error(request, "No file uploaded.")
            return render(request, 'admin_panel/import_admitted_students.html', {'semester': semester})
        
        uploaded_file = request.FILES['file']
        
        try:
            import pandas as pd
            
            # Read file (support CSV and Excel)
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            elif uploaded_file.name.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(uploaded_file)
            else:
                messages.error(request, "File must be CSV or Excel format.")
                return render(request, 'admin_panel/import_admitted_students.html', {'semester': semester})
            
            # Expected columns
            required_cols = {'matricule', 'first_name', 'last_name'}
            optional_cols = {'email', 'program', 'level'}
            
            df_cols = set(df.columns)
            
            if not required_cols.issubset(df_cols):
                messages.error(request, f"Missing required columns: {', '.join(required_cols - df_cols)}")
                return render(request, 'admin_panel/import_admitted_students.html', {'semester': semester})
            
            created_count = 0
            updated_count = 0
            errors = []
            
            with transaction.atomic():
                for index, row in df.iterrows():
                    try:
                        matricule = str(row['matricule']).strip()
                        first_name = str(row['first_name']).strip()
                        last_name = str(row['last_name']).strip()
                        
                        if not matricule or not first_name or not last_name:
                            errors.append(f"Row {index + 2}: Missing required data")
                            continue
                        
                        defaults = {
                            'first_name': first_name,
                            'last_name': last_name,
                            'email': str(row.get('email', '')).strip() if 'email' in df_cols else None,
                            'program': str(row.get('program', '')).strip() if 'program' in df_cols else None,
                            'level': str(row.get('level', '')).strip() if 'level' in df_cols and str(row.get('level', '')).strip() in ['bachelor', 'master', 'phd'] else None,
                        }
                        
                        obj, created = AdmittedStudent.objects.update_or_create(
                            semester=semester,
                            matricule=matricule,
                            defaults=defaults
                        )
                        
                        if created:
                            created_count += 1
                        else:
                            updated_count += 1
                    
                    except Exception as e:
                        errors.append(f"Row {index + 2}: {str(e)}")
            
            if errors:
                error_msg = "\n".join(errors[:10])
                messages.warning(request, f"Import completed with {len(errors)} error(s). First few:\n{error_msg}")
            
            messages.success(request, f"✅ Import complete! {created_count} students created, {updated_count} updated.")
            return redirect('admin_dashboard')
        
        except Exception as e:
            messages.error(request, f"Import failed: {str(e)}")
            return render(request, 'admin_panel/import_admitted_students.html', {'semester': semester})


@login_required
def lookup_admitted_student(request):
    """
    Search across ALL semesters — a student admitted once stays admitted forever.
    Returns JSON with student details.
    """
    matricule = request.GET.get('matricule', '').strip()
    if not matricule:
        return JsonResponse({'found': False, 'message': 'Please enter a matricule'})
    
    # Search across ALL semesters — take most recent record if duplicates
    admitted = AdmittedStudent.objects.filter(
        matricule__iexact=matricule
    ).order_by('-semester__created_at').first()
    
    if admitted:
        return JsonResponse({
            'found': True,
            'matricule': admitted.matricule,
            'first_name': admitted.first_name,
            'last_name': admitted.last_name,
            'email': admitted.email or '',
            'program': admitted.program or '',
            'level': admitted.level or '',
        })
    else:
        return JsonResponse({
            'found': False,
            'message': f'No student found with matricule "{matricule}" in the system'
        })


@login_required
@role_required(['registra'])
def add_student_to_coded_course(request, course_id):
    """
    Add a student (from admitted students database) to a coded course.
    Search across ALL semesters — not just active.
    """
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Method not allowed'})
    
    import json
    data = json.loads(request.body)
    matricule = data.get('matricule', '').strip()
    
    # Search across ALL semesters — not just active
    admitted = AdmittedStudent.objects.filter(
        matricule__iexact=matricule
    ).order_by('-semester__created_at').first()
    
    if not admitted:
        return JsonResponse({
            'success': False,
            'error': f'Student "{matricule}" not found in the admitted students database'
        })
    
    course = get_object_or_404(Course, id=course_id)
    
    # Create Student if not exists — no duplication
    student, _ = Student.objects.update_or_create(
        matricule=admitted.matricule,
        defaults={
            'first_name': admitted.first_name,
            'last_name': admitted.last_name,
            'email': admitted.email or '',
        }
    )
    
    # Check if already enrolled
    if Enrollment.objects.filter(student=student, course=course).exists():
        return JsonResponse({
            'success': False,
            'error': f'{student.get_full_name()} is already enrolled in this course'
        })
    
    # Assign next anonymous code (properly sequenced: 1, 2, 3...)
    max_code = Enrollment.objects.filter(
        course=course
    ).aggregate(Max('anonymous_code'))['anonymous_code__max'] or 0
    next_code = max_code + 1
    
    try:
        enrollment = Enrollment.objects.create(
            student=student,
            course=course,
            semester=course.semester,
            anonymous_code=next_code
        )
        
        # Create associated Grade object for the new enrollment
        from results.models import Grade
        Grade.objects.get_or_create(enrollment=enrollment)
        
        return JsonResponse({
            'success': True,
            'anonymous_code': next_code,
            'student_name': student.get_full_name(),
            'matricule': student.matricule,
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Error adding student: {str(e)}'
        })
